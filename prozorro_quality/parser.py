from __future__ import annotations

import csv
import zipfile
import shutil
import subprocess
from pathlib import Path
from xml.etree import ElementTree

from bs4 import BeautifulSoup
from docx import Document
from openpyxl import load_workbook
from pdfminer.high_level import extract_text as extract_pdf_text
import xlrd

from .config import Settings
from .models import DocumentResult
from .prozorro_api import suffix_for_document


class DocumentParser:
    def __init__(self, settings: Settings):
        self.settings = settings

    def parse(self, document: DocumentResult) -> tuple[DocumentResult, str]:
        if not document.local_path:
            if not document.limitation:
                document.status = "не завантажено"
                document.limitation = "Документ не був завантажений."
            return document, ""

        path = Path(document.local_path)
        suffix = path.suffix.lower() or suffix_for_document(document)
        try:
            if suffix == ".pdf":
                text = extract_pdf_text(str(path)) or ""
            elif suffix == ".docx":
                text = read_docx(path)
            elif suffix == ".doc":
                text = read_legacy_doc(path)
            elif suffix == ".xlsx":
                text = read_xlsx(path)
            elif suffix == ".xls":
                text = read_xls(path)
            elif suffix == ".rtf":
                text = read_rtf(path)
            elif suffix in {".html", ".htm"}:
                text = read_html(path)
            elif suffix in {".txt", ".csv"}:
                text = read_text_or_csv(path, suffix)
            else:
                document.status = "не підтримується"
                document.limitation = (
                    "Формат документа не підтримується MVP-парсером; потрібна ручна перевірка."
                )
                return document, ""
        except Exception as exc:
            document.status = "помилка парсингу"
            document.limitation = f"Не вдалося витягнути текст: {exc}"
            return document, ""

        text = normalize_text(text)
        if len(text) > self.settings.max_text_chars_per_doc:
            text = text[: self.settings.max_text_chars_per_doc]
            document.limitation = (
                "Текст документа обрізано для MVP-аналізу; довгі розділи потребують перевірки."
            )
        document.parsed_chars = len(text)
        if len(text) < 250:
            document.status = "мало тексту"
            document.limitation = (
                document.limitation
                or "У документі мало витягнутого тексту; він може бути сканованим або складним для аналізу."
            )
        else:
            document.status = "опрацьовано"
        return document, text


def read_docx(path: Path) -> str:
    try:
        doc = Document(str(path))
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        try:
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
        except Exception:
            parts.append(read_docx_xml(path))
        return "\n".join(part for part in parts if part)
    except Exception:
        return read_docx_xml(path)


def read_docx_xml(path: Path) -> str:
    parts: list[str] = []
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            if name.startswith("word/media/"):
                continue
            raw = archive.read(name)
            try:
                root = ElementTree.fromstring(raw)
            except ElementTree.ParseError:
                continue
            for node in root.iter():
                if node.tag.endswith("}t") and node.text:
                    parts.append(node.text)
                elif node.tag.endswith("}tab"):
                    parts.append(" ")
                elif node.tag.endswith("}br") or node.tag.endswith("}p"):
                    parts.append("\n")
    return "".join(parts)


def read_legacy_doc(path: Path) -> str:
    if shutil.which("textutil"):
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=20,
        )
        return result.stdout
    if shutil.which("pandoc"):
        result = subprocess.run(
            ["pandoc", "-t", "plain", str(path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=25,
        )
        return result.stdout
    raise RuntimeError("Немає textutil або pandoc для старого формату DOC.")


def read_rtf(path: Path) -> str:
    if shutil.which("textutil"):
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=20,
        )
        return result.stdout
    if shutil.which("pandoc"):
        result = subprocess.run(
            ["pandoc", "-t", "plain", str(path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=20,
        )
        return result.stdout
    raw = path.read_text(encoding="utf-8", errors="ignore")
    return " ".join(part for part in raw.replace("\\par", "\n").split() if not part.startswith("\\"))


def read_xlsx(path: Path) -> str:
    workbook = load_workbook(path, data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"Аркуш: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def read_xls(path: Path) -> str:
    workbook = xlrd.open_workbook(str(path), on_demand=True)
    parts: list[str] = []
    for sheet in workbook.sheets():
        parts.append(f"Аркуш: {sheet.name}")
        for row_idx in range(sheet.nrows):
            values = []
            for col_idx in range(sheet.ncols):
                value = sheet.cell_value(row_idx, col_idx)
                if value not in ("", None):
                    values.append(str(value).strip())
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def read_html(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(raw, "html.parser")
    return soup.get_text("\n")


def read_text_or_csv(path: Path, suffix: str) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    if suffix != ".csv":
        return raw
    rows: list[str] = []
    for row in csv.reader(raw.splitlines()):
        values = [value.strip() for value in row if value.strip()]
        if values:
            rows.append(" | ".join(values))
    return "\n".join(rows)


def normalize_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.replace("\x00", " ").splitlines()]
    return "\n".join(line for line in lines if line)
